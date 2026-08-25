"""Text extraction for student-attached files (chat uploads).

Unlike the admin ingestion pipeline, these files are NOT added to the
knowledge base — the text is used only as context for one question.
"""

import io
import zipfile

import pymupdf
from docx import Document

MAX_WORDS = 8000  # keep the prompt within a sane context size
MAX_PDF_PAGES = 100
MAX_DOCX_FILES = 2000
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024


def _validate_docx_archive(data: bytes) -> None:
    """Reject malformed or highly-expanded DOCX archives before parsing."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            entries = archive.infolist()
            names = {entry.filename for entry in entries}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ValueError("The attached file is not a valid DOCX document.")
            if len(entries) > MAX_DOCX_FILES:
                raise ValueError("The DOCX contains too many internal files.")
            if sum(entry.file_size for entry in entries) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ValueError("The DOCX expands beyond the safe processing limit.")
    except zipfile.BadZipFile as exc:
        raise ValueError("The attached file is not a valid DOCX document.") from exc


def extract_text(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        try:
            doc = pymupdf.open(stream=data, filetype="pdf")
            try:
                if doc.page_count > MAX_PDF_PAGES:
                    raise ValueError(
                        f"PDF has too many pages (max {MAX_PDF_PAGES})."
                    )
                text = "\n\n".join(page.get_text("text") for page in doc)
            finally:
                doc.close()
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError("The attached file is not a readable PDF.") from exc
    elif name.endswith(".docx"):
        _validate_docx_archive(data)
        try:
            docx = Document(io.BytesIO(data))
            parts = [p.text for p in docx.paragraphs if p.text.strip()]
            for table in docx.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text.strip() for cell in row.cells))
            text = "\n\n".join(parts)
        except Exception as exc:
            raise ValueError("The attached file is not a readable DOCX document.") from exc
    elif name.endswith(".txt"):
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError(
            "Unsupported file type. Please attach a PDF, DOCX, or TXT file "
            "(old .doc files are not supported — save as .docx)."
        )

    words = text.split()
    if not words:
        raise ValueError(
            "No readable text found in the file. If it's a scanned document, "
            "the pages are images and can't be read yet."
        )
    truncated = len(words) > MAX_WORDS
    text = " ".join(words[:MAX_WORDS])
    if truncated:
        text += "\n\n[Document truncated — only the first part is shown.]"
    return text
